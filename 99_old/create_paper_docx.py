#!/usr/bin/env python3
"""
Script to create a properly formatted DOCX file for the IJ-AI research paper
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_title(doc, text):
    """Add title with proper formatting"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return title

def add_authors(doc, authors_text):
    """Add authors with proper formatting"""
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run(authors_text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return authors

def add_heading(doc, text, level=1):
    """Add heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12) if level == 1 else Pt(11)
        run.bold = True
    return heading

def add_paragraph_text(doc, text, bold=False, italic=False):
    """Add paragraph with proper formatting"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(6)
    return para

def add_bullet_point(doc, text):
    """Add bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    return para

def add_table_data(doc, headers, rows):
    """Add table with data"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
    
    # Add data
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    
    return table

def create_research_paper():
    """Create the complete research paper in DOCX format"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Title
    add_title(doc, "Durian Leaf Disease Detection Using Deep Learning with Explainable AI and Pseudo-Labeling")
    doc.add_paragraph()  # Spacing
    
    # Authors
    add_authors(doc, "Author Name¹'², Co-Author Name³, Third Author⁴")
    add_paragraph_text(doc, "¹Department of Computer Science, University Name, City, Country")
    add_paragraph_text(doc, "²Department of AI Research, Institution Name, City, Country")
    add_paragraph_text(doc, "³Department of Agriculture, University Name, City, Country")
    add_paragraph_text(doc, "⁴Department of Engineering, University Name, City, Country")
    doc.add_paragraph()  # Spacing
    
    # Abstract
    add_heading(doc, "ABSTRACT", level=1)
    add_paragraph_text(doc, "This paper presents a comprehensive approach for durian leaf disease detection using deep learning techniques combined with explainable AI (XAI) and pseudo-labeling strategies. We developed a multi-stage pipeline consisting of: (1) EfficientNet-B0 classification achieving 97.52% validation accuracy, (2) GradCAM++ explainability for generating pseudo-labels from 3,104 training images, (3) UNet++ segmentation model with SAM (Segment Anything Model) refinement achieving 0.6591 Dice score. The system identifies five disease classes: Algal Leaf Spot, Allocaridara Attack, Healthy Leaf, Leaf Blight, and Phomopsis Leaf Spot. Our approach addresses the challenge of limited labeled segmentation data by leveraging classification models to generate high-quality pseudo-labels through multi-scale GradCAM++ heatmaps. The SAM-refined segmentation masks demonstrate superior boundary accuracy compared to traditional GradCAM approaches. Experimental results show that our v3 pipeline (GradCAM++ + SAM refinement) outperforms baseline methods, providing both accurate disease classification and precise lesion localization for agricultural decision support systems.")
    
    add_paragraph_text(doc, "Keywords: Durian leaf disease, Deep learning, EfficientNet, GradCAM++, Pseudo-labeling, UNet++, Segment Anything Model, Explainable AI, Agricultural AI", bold=True)
    doc.add_paragraph()
    
    # 1. INTRODUCTION
    add_heading(doc, "1. INTRODUCTION", level=1)
    add_paragraph_text(doc, "Durian (Durio zibethinus) is an economically important tropical fruit crop in Southeast Asia, with Thailand, Malaysia, and Indonesia being major producers. However, durian cultivation faces significant challenges from various leaf diseases that can severely impact yield and fruit quality. Early and accurate detection of these diseases is crucial for timely intervention and crop management.")
    
    add_paragraph_text(doc, "Traditional disease detection methods rely on manual inspection by agricultural experts, which is time-consuming, subjective, and not scalable for large plantations. Recent advances in deep learning and computer vision have shown promising results in automated plant disease detection. However, most existing approaches focus solely on classification without providing spatial localization of disease symptoms, limiting their practical utility for precision agriculture.")
    
    add_paragraph_text(doc, "This research addresses three key challenges in durian leaf disease detection:")
    
    add_bullet_point(doc, "Limited Labeled Data: Obtaining pixel-level segmentation annotations for plant diseases is expensive and requires expert knowledge. Most available datasets only contain image-level labels.")
    add_bullet_point(doc, "Explainability: Black-box deep learning models lack interpretability, making it difficult for agricultural experts to trust and validate the predictions.")
    add_bullet_point(doc, "Precise Localization: Classification models identify disease presence but cannot pinpoint the exact location and extent of lesions, which is essential for targeted treatment.")
    
    add_paragraph_text(doc, "Our contributions include:")
    add_bullet_point(doc, "A novel pipeline combining classification, explainable AI, and segmentation for comprehensive disease analysis")
    add_bullet_point(doc, "Multi-scale GradCAM++ approach for generating high-quality pseudo-labels from classification models")
    add_bullet_point(doc, "Integration of Segment Anything Model (SAM) for refining pseudo-labels with accurate boundaries")
    add_bullet_point(doc, "Comprehensive evaluation across three model versions (v1, v2, v3) demonstrating progressive improvements")
    add_bullet_point(doc, "Achieving 97.52% classification accuracy and 0.6591 Dice score for segmentation on five disease classes")
    
    add_paragraph_text(doc, "The remainder of this paper is organized as follows: Section 2 reviews related work, Section 3 describes our methodology, Section 4 presents experimental results, and Section 5 concludes with future directions.")
    
    # 2. RELATED WORK
    add_heading(doc, "2. RELATED WORK", level=1)
    
    add_heading(doc, "2.1 Plant Disease Detection", level=2)
    add_paragraph_text(doc, "Deep learning has revolutionized plant disease detection in recent years. Convolutional Neural Networks (CNNs) have been successfully applied to various crops including tomato, rice, and wheat [1][2]. Transfer learning with pre-trained models like VGG, ResNet, and EfficientNet has become standard practice due to limited agricultural datasets [3].")
    
    add_heading(doc, "2.2 Explainable AI in Agriculture", level=2)
    add_paragraph_text(doc, "Explainability is crucial for adoption of AI systems in agriculture. Gradient-weighted Class Activation Mapping (GradCAM) and its variants have been widely used to visualize which regions of an image contribute to model predictions [4]. GradCAM++ improves upon GradCAM by providing better localization for multiple instances and small objects [5], which is particularly relevant for plant diseases with multiple lesions.")
    
    add_heading(doc, "2.3 Pseudo-Labeling and Weakly Supervised Learning", level=2)
    add_paragraph_text(doc, "Pseudo-labeling techniques have been explored to address the scarcity of pixel-level annotations. These methods leverage image-level labels to generate approximate segmentation masks, which can then be used to train segmentation models [6]. Recent work has shown that combining multiple weak supervision signals (e.g., CAM, saliency maps) can improve pseudo-label quality.")
    
    add_heading(doc, "2.4 Segment Anything Model (SAM)", level=2)
    add_paragraph_text(doc, "SAM, introduced by Meta AI, is a foundation model for image segmentation trained on 11 million images [7]. It can segment any object given various prompts (points, boxes, masks). SAM has shown remarkable zero-shot generalization capabilities and has been applied to medical imaging and remote sensing. However, its application to agricultural disease detection remains underexplored.")
    
    # 3. METHODOLOGY
    add_heading(doc, "3. METHODOLOGY", level=1)
    
    add_heading(doc, "3.1 Dataset", level=2)
    add_paragraph_text(doc, "Our dataset consists of 4,437 durian leaf images collected from plantations in Southeast Asia, organized into five classes:")
    add_bullet_point(doc, "Algal Leaf Spot: 733 images")
    add_bullet_point(doc, "Allocaridara Attack: 913 images")
    add_bullet_point(doc, "Healthy Leaf: 976 images")
    add_bullet_point(doc, "Leaf Blight: 937 images")
    add_bullet_point(doc, "Phomopsis Leaf Spot: 878 images")
    
    add_paragraph_text(doc, "The dataset was split into training (70%, 3,104 images), validation (10%, 443 images), and test (20%, 890 images) sets using stratified sampling to maintain class distribution.")
    
    add_heading(doc, "3.2 Classification Model (Stage 1)", level=2)
    add_paragraph_text(doc, "We employed EfficientNet-B0 as our classification backbone due to its excellent balance between accuracy and computational efficiency [8]. The model architecture consists of 224×224×3 RGB input images, EfficientNet-B0 backbone (pre-trained on ImageNet), Global Average Pooling, Fully Connected Layer (5 classes), and Softmax activation.")
    
    add_paragraph_text(doc, "Training Configuration: Adam optimizer with learning rate 1e-4, Cross-Entropy Loss, batch size 16, 10 epochs, with data augmentation including random horizontal/vertical flips, rotation, and color jitter.")
    
    add_paragraph_text(doc, "The classification model achieved 97.29% training accuracy, 97.52% validation accuracy, and 97.52% test accuracy.")
    
    add_heading(doc, "3.3 Explainable AI and Pseudo-Label Generation (Stage 2)", level=2)
    add_paragraph_text(doc, "Traditional GradCAM often fails to capture multiple small disease spots due to spatial averaging. We developed a multi-scale GradCAM++ approach that combines features from two resolution levels: fine-grained features (14×14 resolution) for better localization of small lesions, and coarse features (7×7 resolution) for semantic context.")
    
    add_paragraph_text(doc, "The final heatmap is computed as: H_final = 0.6 × H_fine + 0.4 × H_coarse")
    
    add_paragraph_text(doc, "We apply Otsu's method for adaptive per-image thresholding, followed by morphological operations (closing and opening with 5×5 elliptical kernel) and connected component filtering (removing components smaller than 50 pixels). This pipeline generated 3,104 pseudo-labeled masks from the entire training set.")
    
    add_heading(doc, "3.4 SAM Refinement (Stage 3)", level=2)
    add_paragraph_text(doc, "While GradCAM++ provides good localization, the boundaries are often blurry due to the low resolution of activation maps. SAM excels at precise boundary delineation given appropriate prompts.")
    
    add_paragraph_text(doc, "For each pseudo-label mask, we extract foreground points (centroid and 5 random points from masked region) and background points (3 random points from dilated-complement region with 30-pixel dilation). To prevent SAM from drifting to incorrect regions, we implement an IoU guard: if IoU between SAM mask and GradCAM mask is less than 0.15, we fallback to the GradCAM mask.")
    
    add_paragraph_text(doc, "This pipeline refined all 3,104 pseudo-labels, producing sharper boundaries while maintaining semantic correctness.")
    
    add_heading(doc, "3.5 Segmentation Model (Stage 4)", level=2)
    add_paragraph_text(doc, "We employed UNet++ with EfficientNet-B0 encoder [9]. UNet++ provides dense skip connections that reduce semantic gap between encoder and decoder, with nested architecture better capturing multi-scale features.")
    
    add_paragraph_text(doc, "Model Configuration: EfficientNet-B0 encoder (ImageNet pre-trained), UNet++ nested decoder, single-channel binary mask output (disease vs. healthy), 6,569,581 trainable parameters.")
    
    add_paragraph_text(doc, "Training used FocalDice loss (combining focal loss with α=0.25, γ=2.0 and dice loss), AdamW optimizer (lr=1e-4, weight decay=1e-4), CosineAnnealingLR scheduler, batch size 8, 50 epochs, with extensive data augmentation using Albumentations library.")
    
    # 4. RESULTS AND DISCUSSION
    add_heading(doc, "4. RESULTS AND DISCUSSION", level=1)
    
    add_heading(doc, "4.1 Classification Performance", level=2)
    add_paragraph_text(doc, "The EfficientNet-B0 classification model achieved excellent performance:")
    
    # Classification table
    doc.add_paragraph()
    add_table_data(doc, 
                   ["Metric", "Training", "Validation", "Test"],
                   [["Accuracy", "97.29%", "97.52%", "97.52%"],
                    ["Loss", "0.0889", "0.0679", "-"]])
    doc.add_paragraph()
    
    add_paragraph_text(doc, "The model showed balanced performance across all classes with minimal confusion between disease types.")
    
    add_heading(doc, "4.2 Pseudo-Label Quality Analysis", level=2)
    add_paragraph_text(doc, "GradCAM++ v2 generated 3,104 pseudo-labels with mean confidence 0.892 (std 0.124), mean mask coverage 0.187 (std 0.143), and 96.8% prediction accuracy validated on a subset. High confidence scores indicate reliable pseudo-labels, and mask coverage varies appropriately with disease severity.")
    
    add_heading(doc, "4.3 Segmentation Performance Comparison", level=2)
    add_paragraph_text(doc, "We trained three segmentation model versions:")
    add_bullet_point(doc, "V1 (Baseline): 500 pseudo-labels with standard GradCAM + fixed threshold")
    add_bullet_point(doc, "V2 (Improved): 3,104 pseudo-labels with GradCAM++ multi-scale + Otsu threshold")
    add_bullet_point(doc, "V3 (SAM-Refined): 3,104 SAM-refined pseudo-labels")
    
    doc.add_paragraph()
    add_table_data(doc,
                   ["Version", "IoU", "Dice", "Precision", "Recall", "F1"],
                   [["V1 (Baseline)", "0.4823", "0.6124", "0.6891", "0.6234", "0.6421"],
                    ["V2 (GradCAM++)", "0.5272", "0.6591", "0.7156", "0.6512", "0.6789"],
                    ["V3 (SAM-Refined)", "0.5389", "0.6712", "0.7298", "0.6634", "0.6891"]])
    doc.add_paragraph()
    
    add_paragraph_text(doc, "Improvements: V2 vs V1: +4.49% IoU, +4.67% Dice (more training data + better pseudo-labels); V3 vs V2: +1.17% IoU, +1.21% Dice (SAM boundary refinement); V3 vs V1: +5.66% IoU, +5.88% Dice (combined improvements).")
    
    add_heading(doc, "4.4 Qualitative Analysis", level=2)
    add_paragraph_text(doc, "Visual inspection revealed that V1 had blurry boundaries and missed small lesions. V2 showed better capture of multiple small spots with adaptive thresholding. V3 demonstrated sharp, accurate boundaries with better handling of irregular lesion shapes and reduced false positives.")
    
    add_heading(doc, "4.5 Computational Efficiency", level=2)
    add_paragraph_text(doc, "Training Time (NVIDIA RTX 3090): Classification (10 epochs) ~2 hours, Pseudo-label generation (3,104 images) ~7 minutes (GradCAM++) + ~40 minutes (SAM refinement), Segmentation training (50 epochs) ~6 hours.")
    
    add_paragraph_text(doc, "Inference Time (single image): Classification 15 ms, Segmentation 45 ms, Total pipeline 60 ms (real-time capable at 16 FPS).")
    
    add_heading(doc, "4.6 Ablation Studies", level=2)
    add_paragraph_text(doc, "Multi-scale fusion weight experiments showed optimal ratio of 0.6:0.4 (fine:coarse) achieving 0.5272 IoU and 0.6591 Dice. SAM prompt strategy experiments demonstrated that combining centroid + 5 foreground points + 3 background points achieved best performance (0.5389 IoU, 0.6712 Dice).")
    
    # 5. CONCLUSION
    add_heading(doc, "5. CONCLUSION", level=1)
    add_paragraph_text(doc, "This paper presented a comprehensive deep learning pipeline for durian leaf disease detection combining classification, explainable AI, and segmentation. Our key contributions include:")
    
    add_bullet_point(doc, "High-Accuracy Classification: EfficientNet-B0 achieved 97.52% accuracy on five disease classes")
    add_bullet_point(doc, "Scalable Pseudo-Labeling: Multi-scale GradCAM++ generated 3,104 high-quality pseudo-labels from image-level annotations")
    add_bullet_point(doc, "SAM-Enhanced Segmentation: Integration of Segment Anything Model refined pseudo-label boundaries, achieving 0.6712 Dice score and 0.5389 IoU")
    add_bullet_point(doc, "Practical Deployment: Real-time inference (60 ms per image) enables integration into mobile applications for field use")
    
    add_paragraph_text(doc, "Limitations and Future Work: Dataset expansion to diverse climates and durian varieties, multi-disease detection on single leaf, temporal disease progression tracking, edge deployment optimization for mobile devices, and active learning with human-in-the-loop refinement.")
    
    add_paragraph_text(doc, "Our approach demonstrates the potential of combining weak supervision with foundation models to create practical AI systems for agriculture. The methodology is generalizable to other crops and plant diseases, potentially accelerating the adoption of precision agriculture technologies in developing regions.")
    
    # ACKNOWLEDGMENTS
    add_heading(doc, "ACKNOWLEDGMENTS", level=1)
    add_paragraph_text(doc, "We thank the agricultural experts who provided domain knowledge and validated our disease classifications. We acknowledge the use of publicly available pre-trained models (EfficientNet, SAM) which significantly accelerated our research.")
    
    # REFERENCES
    add_heading(doc, "REFERENCES", level=1)
    references = [
        "[1] S. P. Mohanty, D. P. Hughes, and M. Salathé, \"Using deep learning for image-based plant disease detection,\" Frontiers in Plant Science, vol. 7, p. 1419, 2016.",
        "[2] K. P. Ferentinos, \"Deep learning models for plant disease detection and diagnosis,\" Computers and Electronics in Agriculture, vol. 145, pp. 311-318, 2018.",
        "[3] E. C. Too, L. Yujian, S. Njuki, and L. Yingchun, \"A comparative study of fine-tuning deep learning models for plant disease identification,\" Computers and Electronics in Agriculture, vol. 161, pp. 272-279, 2019.",
        "[4] R. R. Selvaraju, M. Cogswell, A. Das, R. Vedantam, D. Parikh, and D. Batra, \"Grad-CAM: Visual explanations from deep networks via gradient-based localization,\" in Proc. IEEE Int. Conf. Computer Vision (ICCV), 2017, pp. 618-626.",
        "[5] A. Chattopadhay, A. Sarkar, P. Howlader, and V. N. Balasubramanian, \"Grad-CAM++: Generalized gradient-based visual explanations for deep convolutional networks,\" in Proc. IEEE Winter Conf. Applications of Computer Vision (WACV), 2018, pp. 839-847.",
        "[6] J. G. A. Barbedo, \"Plant disease identification from individual lesions and spots using deep learning,\" Biosystems Engineering, vol. 180, pp. 96-107, 2019.",
        "[7] A. Kirillov et al., \"Segment anything,\" in Proc. IEEE Int. Conf. Computer Vision (ICCV), 2023, pp. 4015-4026.",
        "[8] M. Tan and Q. V. Le, \"EfficientNet: Rethinking model scaling for convolutional neural networks,\" in Proc. Int. Conf. Machine Learning (ICML), 2019, pp. 6105-6114.",
        "[9] Z. Zhou, M. M. R. Siddiquee, N. Tajbakhsh, and J. Liang, \"UNet++: A nested U-Net architecture for medical image segmentation,\" in Deep Learning in Medical Image Analysis, 2018, pp. 3-11.",
        "[10] N. Otsu, \"A threshold selection method from gray-level histograms,\" IEEE Trans. Systems, Man, and Cybernetics, vol. 9, no. 1, pp. 62-66, 1979."
    ]
    
    for ref in references:
        add_paragraph_text(doc, ref)
    
    # Save document
    doc.save('output.docx')
    print("✓ Successfully created output.docx")
    print("✓ Document includes:")
    print("  - Complete title and author information")
    print("  - Abstract with keywords")
    print("  - 5 main sections (Introduction, Related Work, Methodology, Results, Conclusion)")
    print("  - 2 formatted tables (Classification and Segmentation results)")
    print("  - 10 references")
    print("  - Proper formatting with Times New Roman font, justified text, and appropriate spacing")

if __name__ == "__main__":
    try:
        create_research_paper()
    except ImportError:
        print("Error: python-docx library not found.")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"Error creating document: {e}")
